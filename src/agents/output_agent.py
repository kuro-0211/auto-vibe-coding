import os
import io
import re
from datetime import datetime
from dotenv import load_dotenv
load_dotenv("/app/.env")
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage
from utils.logger import pipeline_logger

OUTPUT_DIR = "/app/data/outputs"


def run_output(state: dict) -> str:
    pipeline_logger.log_step("Output Agent", "running")

    llm = ChatOllama(
        model=os.getenv("GEMMA_MODEL", "gemma3:4b"),
        base_url=os.getenv("OLLAMA_BASE_URL", "http://ollama:11434"),
        temperature=0.3
    )

    research = state.get("research_result", "")
    code = state.get("code_result", "")
    execution = state.get("execution_result")
    retry_count = state.get("retry_count", 0)

    if execution and isinstance(execution, dict):
        if execution.get("success"):
            exec_summary = f"실행 성공\n{execution.get('output', '')}"
        else:
            exec_summary = f"실행 실패 ({retry_count}회 시도)\n{execution.get('error', '')}"
    else:
        exec_summary = "코드 실행 없음"

    prompt = f"""
다음 내용을 바탕으로 최종 결과 문서를 작성해주세요.

## 리서치 결과
{research}

## 생성된 코드
{code if code else "코드 생성 없음"}

## 실행 결과
{exec_summary}

## 출력 형식
# 결과 요약
(핵심 내용 3~5줄)

## 리서치 내용
(정리된 내용)

## 코드
(생성된 코드, 있는 경우)

## 실행 결과
(실행 결과, 있는 경우)

한국어로 작성하세요.
"""

    response = llm.invoke([HumanMessage(content=prompt)])

    pipeline_logger.log_llm(
        model="gemma3:4b",
        prompt=prompt,
        response=response.content,
        tokens=0
    )
    pipeline_logger.log_step("Output Agent", "done", output_data=response.content)

    return response.content


def _ensure_output_dir() -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    return OUTPUT_DIR


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


_INLINE_PATTERNS = [
    (re.compile(r"\*\*(.+?)\*\*"), r"\1"),   # **bold**
    (re.compile(r"__(.+?)__"),     r"\1"),   # __bold__
    (re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)"), r"\1"),  # *italic*
    (re.compile(r"(?<!_)_([^_\n]+?)_(?!_)"),     r"\1"),  # _italic_
    (re.compile(r"~~(.+?)~~"),     r"\1"),   # ~~strike~~
    (re.compile(r"`([^`\n]+?)`"),  r"\1"),   # `inline code`
    (re.compile(r"!\[([^\]]*)\]\([^)]+\)"), r"[이미지: \1]"),  # ![alt](url)
    (re.compile(r"\[([^\]]+)\]\(([^)]+)\)"), r"\1 (\2)"),     # [text](url)
]


def _strip_inline_md(text: str) -> str:
    """본문 한 줄에서 마크다운 inline 마커 제거."""
    for pat, repl in _INLINE_PATTERNS:
        text = pat.sub(repl, text)
    return text


def _parse_blocks(content: str) -> list[tuple[str, str]]:
    """마크다운을 (block_type, text) 시퀀스로 변환.

    block_type ∈ {h1, h2, h3, code, bullet, num, blank, p}
    text는 코드블록이면 여러 줄(\\n 포함), 그 외엔 한 줄.
    """
    blocks: list[tuple[str, str]] = []
    in_code = False
    code_buf: list[str] = []

    for raw in content.splitlines():
        line = raw.rstrip()
        # 코드블록 진입/이탈
        if line.lstrip().startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code_buf)))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        stripped = line.strip()
        if stripped.startswith("# "):
            blocks.append(("h1", _strip_inline_md(stripped[2:])))
        elif stripped.startswith("## "):
            blocks.append(("h2", _strip_inline_md(stripped[3:])))
        elif stripped.startswith("### "):
            blocks.append(("h3", _strip_inline_md(stripped[4:])))
        elif stripped.startswith(("- ", "* ", "+ ")):
            blocks.append(("bullet", _strip_inline_md(stripped[2:])))
        elif re.match(r"^\d+\.\s+", stripped):
            # 원본 숫자 보존 ("1. 텍스트")
            blocks.append(("num", _strip_inline_md(stripped)))
        elif stripped == "":
            blocks.append(("blank", ""))
        else:
            blocks.append(("p", _strip_inline_md(line)))

    if in_code and code_buf:
        blocks.append(("code", "\n".join(code_buf)))

    return blocks


def export_to_pdf(content: str, code: str = "", filename: str | None = None) -> bytes:
    """마크다운 텍스트를 PDF 바이트로 변환."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # 한글 폰트 + 모노스페이스 폰트 등록
    body_path = next((p for p in [
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ] if os.path.exists(p)), None)
    mono_path = next((p for p in [
        "/usr/share/fonts/truetype/nanum/NanumGothicCoding.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
    ] if os.path.exists(p)), None)

    if body_path:
        pdf.add_font("Body", "", body_path)
        body_font = "Body"
    else:
        body_font = "Helvetica"
    if mono_path:
        pdf.add_font("Mono", "", mono_path)
        mono_font = "Mono"
    else:
        mono_font = body_font

    pdf.set_font(body_font, size=11)
    w = pdf.epw

    for kind, text in _parse_blocks(content):
        if kind == "h1":
            pdf.ln(2)
            pdf.set_font(body_font, size=18)
            pdf.multi_cell(w=w, h=10, text=text)
            pdf.ln(3)
        elif kind == "h2":
            pdf.ln(2)
            pdf.set_font(body_font, size=14)
            pdf.multi_cell(w=w, h=8, text=text)
            pdf.ln(2)
        elif kind == "h3":
            pdf.ln(1)
            pdf.set_font(body_font, size=12)
            pdf.multi_cell(w=w, h=7, text=text)
            pdf.ln(1)
        elif kind == "bullet":
            pdf.set_font(body_font, size=11)
            pdf.multi_cell(w=w, h=6, text="• " + text)
        elif kind == "num":
            pdf.set_font(body_font, size=11)
            pdf.multi_cell(w=w, h=6, text=text)
        elif kind == "code":
            pdf.ln(1)
            # 코드블록 회색 박스 효과: 흰 줄 사이 띄움
            pdf.set_font(mono_font, size=9)
            for cl in text.splitlines() or [""]:
                pdf.multi_cell(w=w, h=5, text=cl if cl else " ")
            pdf.ln(1)
            pdf.set_font(body_font, size=11)
        elif kind == "blank":
            pdf.ln(3)
        else:  # p
            pdf.set_font(body_font, size=11)
            pdf.multi_cell(w=w, h=6, text=text if text else " ")

    if code:
        pdf.add_page()
        pdf.set_font(body_font, size=14)
        pdf.multi_cell(w=w, h=9, text="생성된 코드")
        pdf.ln(3)
        pdf.set_font(mono_font, size=9)
        for cl in code.splitlines():
            pdf.multi_cell(w=w, h=5, text=cl if cl else " ")

    output = pdf.output(dest="S")
    if isinstance(output, str):
        return output.encode("latin-1")
    return bytes(output)


def _set_kor_font(run, name: str = "Malgun Gothic") -> None:
    """python-docx run에 한글(East Asia) 폰트 명시.

    Word는 라틴/동아시아 폰트가 분리되어 있어 eastAsia 속성을 따로 지정해야
    한글이 의도한 폰트로 렌더링된다.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"),     name)
    rFonts.set(qn("w:hAnsi"),     name)
    rFonts.set(qn("w:eastAsia"),  name)
    rFonts.set(qn("w:cs"),        name)


def export_to_docx(content: str, code: str = "") -> bytes:
    """마크다운 텍스트를 Word(docx) 바이트로 변환."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # 기본 스타일에 한글 폰트 설정 (모든 paragraph에 영향)
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    def add_p(text: str, *, bold: bool = False, mono: bool = False, size: int | None = None,
              color: tuple[int, int, int] | None = None, style_name: str | None = None):
        para = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        run = para.add_run(text)
        if mono:
            run.font.name = "Consolas"
            from docx.oxml import OxmlElement
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        else:
            _set_kor_font(run)
        if size:
            run.font.size = Pt(size)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)
        return para

    for kind, text in _parse_blocks(content):
        if kind == "h1":
            h = doc.add_heading(level=1)
            r = h.add_run(text)
            _set_kor_font(r)
        elif kind == "h2":
            h = doc.add_heading(level=2)
            r = h.add_run(text)
            _set_kor_font(r)
        elif kind == "h3":
            h = doc.add_heading(level=3)
            r = h.add_run(text)
            _set_kor_font(r)
        elif kind == "bullet":
            add_p(text, style_name="List Bullet")
        elif kind == "num":
            # Word "List Number" 스타일이 자동으로 숫자 마커를 부여하므로
            # 본문에서 "1." 같은 접두 숫자는 제거하고 텍스트만 넘김
            clean = re.sub(r"^\d+\.\s+", "", text)
            add_p(clean, style_name="List Number")
        elif kind == "code":
            for cl in text.splitlines() or [""]:
                add_p(cl, mono=True, size=9)
        elif kind == "blank":
            doc.add_paragraph()
        else:
            add_p(text)

    if code:
        h = doc.add_heading(level=2)
        r = h.add_run("생성된 코드")
        _set_kor_font(r)
        for cl in code.splitlines():
            add_p(cl, mono=True, size=9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_output_file(content: str, code: str = "", fmt: str = "pdf") -> str:
    """파일로 저장하고 경로 반환 (이메일 첨부용)."""
    _ensure_output_dir()
    fmt = fmt.lower()
    ts = _timestamp()

    if fmt == "pdf":
        data = export_to_pdf(content, code)
        path = os.path.join(OUTPUT_DIR, f"aria_result_{ts}.pdf")
    elif fmt in ("docx", "word"):
        data = export_to_docx(content, code)
        path = os.path.join(OUTPUT_DIR, f"aria_result_{ts}.docx")
    elif fmt in ("md", "markdown"):
        data = content.encode("utf-8")
        path = os.path.join(OUTPUT_DIR, f"aria_result_{ts}.md")
    else:
        raise ValueError(f"지원하지 않는 형식: {fmt}")

    with open(path, "wb") as f:
        f.write(data)
    return path
